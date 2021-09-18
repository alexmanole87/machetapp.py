# import interfata_grafica_macheta as igm
import os

from docx import Document



def activitati(instit):
    if instit == 'Administrația Domeniului Public București-Sector I':
        return 'Măturat și golit coșuri de gunoi în parcuri, greblat manual, curățenie peluze, salubrizare parcuri, deszăpezire'
    elif instit == 'Administrația Domeniului Public București-Sector II':
        return 'Măturat și golit coșuri de gunoi în parcuri, greblat manual, curățenie peluze, salubrizare parcuri, deszăpezire'
    elif instit == 'Administrația Domeniului Public București-Sector III':
        return 'Măturat și golit coșuri de gunoi în parcuri, greblat manual, curățenie peluze, salubrizare parcuri, dezăpezire'
    elif instit == 'Administrația Domeniului Public și Dezvoltare Urbană Sector 6':
        return 'Udare spații verzi, întreținere gazon, tuns gard viu, plantat material dendrologic, lucrări de dezăpezire, întreținere gazon'
    elif instit == 'Administrația Lacuri, Parcuri și Agrement București':
        return 'Întreținerea parcurilor aflate în administrarea Primăriei Capitalei '
    elif instit == 'Direcția Piețe și Gestionare Activități Comerciale Sector IV':
        return 'Activități de salubrizare și întreținere în piețele sectorului 4'
    elif instit == 'SC Amenajare Edilitară și Salubrizare SA Sector V':
        return 'Măturat și golit coșuri de gunoi în parcuri, greblat manual, curățenie peluze, salubrizare parcuri, dezăpezire'
    elif instit == 'Administrația Domeniului Public și Dezvoltare Urbană Sector VI':
        return 'Udare spații verzi, întreținere gazon, tuns gard viu, plantat material dendrologic, lucrări de dezăpezire, întreținere gazon'
    elif instit == 'Poliția Locală a Sectorului I':
        return 'Asigurarea curățeniei la boxele în care sunt adăpostite animalele fără stăpân, distribuirea de hrană și apă pentru animalele fără stăpân din cadrul adăpostului, săpat, greblat, curățenie alei și peluze etc.'
    elif instit == 'Fundația pentru Promovarea Sancțiunilor Comunitare':
        return 'Dezmembrarea manuală a deșeurilor de echipamente electrice și electronice, activități de colectare și sortare a materialelor reciclabile (hârtie, carton, ambalaje plastice, ambalaje metalice, ambalaje din sticlă) și alte activități de organizare, întreținere și curățenie.  '
    elif instit == 'Primăria Municipiului București - Administrația Cimitirelor și Crematoriilor Umane':
        return 'Activități de întreținere'
    elif instit == 'Arhiepiscopia Bucureștilor':
        return 'Activități de întreținere în baza Hotărârii Permanenței Consiliului Eparhial 3456/2015'
    elif instit == 'Direcția Generală de Asistență Socială și Protecție a Copilului Sector I':
        return 'In funcție de rezultatele evaluării vocaționale și a abilităților, de istoricul profesional și medical al persoanelor sancționate: activități de întreținere și curățenie a spațiilor interioare și exterioare, lucrări de salubrizare și colectare selectivă a gunoiului, lucrări de zugrăvire și vopsitorie, dezinfecția și curățenia cărucioarelor și a celorlalte obiecte care ajută beneficiarii la deplasare,  activități pregătitoare în activitatea efectivă a specialiștilor cu beneficiarii serviciului, activități specifice spălătoriilor auto și textilă, tipografiei, lucrului cu ceramică, croitorie ș.a. (activități specifice desfășurate în cadrul întreprinderii sociale, „Nazarcea Grup”)'
    elif instit == 'Direcția Generală de Asistență Socială și Protecție a Copilului Sector II':
        return 'Activități de întreținere și curățenie în cadrul instituțiilor aflate în subordinea direcției'
    elif instit == 'Direcția Generală de Asistență Socială și Protecție a Copilului Sector III':
        return 'Activități de întreținere și curățenie în cadrul instituțiilor aflate în subordinea direcției'
    elif instit == 'Direcția Generală de Asistență Socială și Protecție a Copilului Sector IV':
        return 'Activități de întreținere și curățenie în cadrul instituțiilor aflate în subordinea direcției'
    elif instit == 'Direcția Generală de Asistență Socială și Protecție a Copilului Sector V':
        return 'Activități de întreținere și curățenie în cadrul instituțiilor aflate în subordinea direcției'
    elif instit == 'Direcția Generală de Asistență Socială și Protecție a Protecția a Copilului Sector VI':
        return 'Activități de întreținere și curățenie în cadrul instituțiilor aflate în subordinea direcției'
    elif instit == 'La alegerea SP':
        return 'La alegerea SP'
    else:
        return ''


class document_supraveghere():
    def __init__(self, dsSuprav, nume, numeP, datan, adresa, cnp, studii, infr1, infr2, infr3, dsinst, nrSentinta,
                 instanta, defin, pedeapsa, ts, indicativ, dataVenire, oraVenire, sector, sectie_pol, data_svp_i,
                 data_svp_f,
                 obligatii, mfc_h, mfc_z, mfc_instit_1_n, mfc_instit_1_a, mfc_instit_2_n, mfc_instit_2_a):
        self.mfc_instit_2_a = mfc_instit_2_a
        self.mfc_instit_2_n = mfc_instit_2_n
        self.mfc_instit_1_a = mfc_instit_1_a
        self.mfc_instit_1_n = mfc_instit_1_n
        self.mfc_h = mfc_h
        self.mfc_z = mfc_z
        self.obligatii = obligatii
        self.data_svp_f = data_svp_f
        self.data_svp_i = data_svp_i
        self.sector = sector
        self.sectie_pol = sectie_pol
        self.dsSuprav = dsSuprav
        self.nume = nume
        self.numeP = numeP
        self.datan = datan
        self.adresa = adresa
        self.cnp = cnp
        self.studii = studii
        self.infr1 = infr1
        self.infr2 = infr2
        self.infr3 = infr3
        self.dsinst = dsinst
        self.nrSentinta = nrSentinta
        self.instanta = instanta
        self.defin = defin
        self.pedeapsa = pedeapsa
        self.ts = ts
        self.indicativ = indicativ
        self.dataVenire = dataVenire
        self.oraVenire = oraVenire

    def pp(self, dir):
        def paragraphs_to_be_added():
            dir = os.getcwd() + "\Macheta Dosar\General\Prima pagina dosar macheta .docx"
            doc_to_be_chanched = Document(dir)
            p0 = str(doc_to_be_chanched.paragraphs[0:6])
            nume_inc = str(doc_to_be_chanched.paragraphs[6].text.replace('(nume)', f"{self.nume}"))
            nume_parinti = str(doc_to_be_chanched.paragraphs[6].text.replace('(prenume parinti)', f"{self.numeP}"))
            data_nasterii = str(
                doc_to_be_chanched.paragraphs[7].text.replace('(data si locul nasterii)', f"{self.datan}"))
            adresa = str(doc_to_be_chanched.paragraphs[8].text.replace('(adresa)', f"{self.adresa}"))
            cnp = str(doc_to_be_chanched.paragraphs[9].text.replace('(CNP)', f"{self.cnp}"))
            studii = str(doc_to_be_chanched.paragraphs[10].text.replace('(nrClase)', f"{self.studii}"))
            p11 = str(doc_to_be_chanched.paragraphs[11:15])
            infr = str(doc_to_be_chanched.paragraphs[15].text.replace('(infracțiune)',
                                                                      f"{self.infr1} + ',' +{self.infr2} + ',' + {self.infr3}"))
            dsInst = str(doc_to_be_chanched.paragraphs[17].text.replace('(dosarInstanță)', f"{self.dsinst}"))
            nrSent = str(doc_to_be_chanched.paragraphs[18].text.replace('(nrSentință)', f"{self.nrSentinta}"))
            numeInst = str(doc_to_be_chanched.paragraphs[18].text.replace('(instanta)', f"{self.instanta}"))
            defSent = str(doc_to_be_chanched.paragraphs[19].text.replace('(def)', f"{self.defin}"))
            pedeapsa = str(doc_to_be_chanched.paragraphs[21].text.replace('(pedeapsa)', f"{self.pedeapsa}"))
            numeTs = str(doc_to_be_chanched.paragraphs[21].text.replace('(termen de supraveghere)', f"{self.ts}"))
            p22 = str(doc_to_be_chanched.paragraphs[22])
            dataStart = str(doc_to_be_chanched.paragraphs[23].text.replace('(data start)', "{self.dataStart}"))
            dataStop = str(doc_to_be_chanched.paragraphs[23].text.replace('(data stop)', "{dataStop}"))
            p_fin = str(doc_to_be_chanched.paragraphs[24:])

            return p0, nume_inc, nume_parinti, data_nasterii, adresa, cnp, studii, p11, infr, dsInst, nrSent, numeInst, defSent, pedeapsa, numeTs, p22, dataStart, dataStop, p_fin

        prima_pagina = Document()
        list_items = paragraphs_to_be_added()

        for e in list_items:
            index = list_items.index(e)
            prima_pagina.add_paragraph(list_items[index])

        prima_pagina.save(dir)
        return prima_pagina

    def confirmare(self, dir):

        # with open(os.getcwd() + '\\Macheta Dosar\\General\\Confirmare de Primire macheta.docx', 'r+',encoding="utf8") as f:
        #     text = f.read()
        #     text = re.sub('\(nume\)', {self.nume}, text)
        #     f.seek(0)
        #     f.write(text)
        #     f.truncate()

        document = Document(os.getcwd() + '\\Macheta Dosar\\General\\Adresa punere executare macheta.docx')

        for paragraph in document.paragraphs:
            if '-nume-' in paragraph.text:
                paragraph.text.replace('-nume-', self.nume)
                print('a fost gasit cuvantul')
                print(paragraph.text)

        document.save(dir)

        return document

    def macheta_lista(self, dir):
        text = f"""(sector)|{self.sector} 
(nrSecție)|{self.sectie_pol}
(dataVenire)|{self.dataVenire}
(oraVenire)|{self.oraVenire}
(nrDosarSuprv)|{self.dsSuprav}
(nume)|{self.nume}
(prenume parinti)|{self.numeP}
(data si locul nasterii)|{self.datan}
(adresa)| {self.adresa}
(CNP)| {self.cnp}
(nrClase)|{self.studii}
(infracțiune)|{self.infr1}\\ {self.infr2}\\{self.infr3}
(dosarInstanță)|{self.dsinst}
(nrSentință)|{self.nrSentinta}
(instanta)| {self.instanta} 
(def)|{self.defin}
(pedeapsa)| {self.pedeapsa}
(indicativ) | {self.indicativ}
(termen de supraveghere)|{self.ts}
(data start)| {self.data_svp_i}
(data stop)| {self.data_svp_f}
(obligatii)|{self.obligatii}
(durată muncă - z)|{self.mfc_z}
(durată muncă - h)|{self.mfc_h}
(instituție 1)|{self.mfc_instit_1_n}
(activități 1)|{self.mfc_instit_1_a}
(instituție 2)|{self.mfc_instit_2_n}
(activități 2)|{self.mfc_instit_2_a}
(SP)|
(adresă_nouă)|
(instanta 1)|
(dosarInstanță 1)|
(nrSentință 1)|
(def 1)|
(termen de supraveghere 1)|
(dosarInstanță 2)|
(nrSentință 2)|
(def 2)|
(termen de supraveghere 2)|
(dosarInstanță 3)|
(nrSentință 3)|
(def 3)|
(termen de supraveghere 3)|
"""
        macheta_lista_completata = Document()
        macheta_lista_completata.add_paragraph(text)
        macheta_lista_completata.save(dir)

        return macheta_lista_completata

# def clickDreaptaOptiuni():
#     menu_cldr = QMenu()
#
#     incarca =menu_cldr.addAction('Încarcă document')
#     reg_ex = menu_cldr.addAction('RE')
#     nlp = menu_cldr.addAction('NLP')
#
#     action = menu_cldr.exec_(mapToGlobal(event.pos()))
